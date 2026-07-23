#!/usr/bin/env python3
"""Render a tailored CV by applying a JSON edit-plan to the base .docx, preserving formatting.

The base (profile/cv-base.docx) is the source of truth. We NEVER rewrite the whole
document — we edit in place so fonts, spacing, list numbering, and live hyperlinks survive untouched:
  * summary  -> replace the text of the single summary paragraph
  * skills   -> per line: keep the bold label run, swap the plain value run; supports reorder/add/drop
  * bullets  -> per experience section: replace the set of list paragraphs (clone the section's own
                bullet paragraph as the template so numbering/style are inherited)

Sections are matched by anchor text (Heading-1 names + bold job-title headers), never by index, so the
script keeps working if the base is edited later.

Usage:
  render_cv.py --base <base.docx> --plan <plan.json> --out <out.docx> [--pdf]

Plan schema (every key optional — omit to leave that part of the base unchanged):
  {
    "title": "Senior Full-Stack Engineer (AI-Focused)",   # the subtitle under the name
    "tagline": "US citizen | Available for remote contract | Will relocate",  # line under contacts
    "summary": "full replacement summary text",
    "skills": [                      # if present, REPLACES the whole skills block in this order
      {"label": "Frontend", "value": "React, TypeScript, ..."},
      ...
    ],
    "experience": [                  # only sections you list are touched; others stay as-is
      # "title" (optional) overrides that entry's bold header line (e.g. add/drop "(Founder)")
      {"match": "Reazy LLC", "title": "Founder & Full-Stack Developer — Reazy LLC",
       "bullets": ["bullet 1", "bullet 2", ...]}
    ]
  }
"""
from __future__ import annotations

import argparse
import copy
import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HEADINGS = ("TECHNICAL SKILLS", "EXPERIENCE", "EDUCATION")


def _text(p: Paragraph) -> str:
    return p.text.strip()


def _is_heading(p: Paragraph) -> bool:
    return p.style.name.startswith("Heading")


def _is_bullet(p: Paragraph) -> bool:
    """True if the paragraph is a numbered/bulleted list item (has w:numPr)."""
    ppr = p._p.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def _paras(doc) -> list[Paragraph]:
    return doc.paragraphs


def _remove(p: Paragraph) -> None:
    p._p.getparent().remove(p._p)


def _set_single_run_text(p: Paragraph, text: str) -> None:
    """Set a paragraph's visible text to `text`, keeping the first run's formatting and dropping the rest.

    Only safe for paragraphs whose meaning lives in one content run (summary, plain bullets). Never call
    this on paragraphs that carry hyperlinks (contact/date lines) — those are never edited by this script.
    """
    runs = p.runs
    if not runs:
        p.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._r.getparent().remove(r._r)


# ----------------------------------------------------------------------------- header (title / tagline)

def _header_paras(doc) -> list[Paragraph]:
    """Non-empty paragraphs of the name/title/contact block (everything before the first Heading-1).

    In order these are: [name, title, contact-line (hyperlinks), tagline, summary]. The tagline is
    optional in the base; the contact line is the one carrying live hyperlinks.
    """
    out = []
    for p in _paras(doc):
        if _is_heading(p):
            break
        if _text(p):
            out.append(p)
    return out


def apply_title(doc, title: str) -> None:
    """The title/subtitle is the second non-empty header paragraph (just under the name). No hyperlinks."""
    header = _header_paras(doc)
    if len(header) < 2:
        raise RuntimeError("could not locate the title paragraph")
    _set_single_run_text(header[1], title)


def apply_tagline(doc, tagline: str) -> None:
    """The tagline is the plain line right after the contact (hyperlink) line, and is not the summary.

    The summary is the last header paragraph before the first heading; the contact line is the one with
    live hyperlinks. The tagline sits between them — editing it is safe (no hyperlinks of its own).
    """
    header = _header_paras(doc)
    if len(header) < 2:
        raise RuntimeError("could not locate the tagline paragraph")
    summary = header[-1]
    contact_idx = next((i for i, p in enumerate(header) if _has_hyperlink(p)), None)
    target = None
    if contact_idx is not None and contact_idx + 1 < len(header):
        cand = header[contact_idx + 1]
        if cand is not summary:
            target = cand
    if target is None:
        raise RuntimeError("base has no distinct tagline line to edit")
    _set_single_run_text(target, tagline)


# ----------------------------------------------------------------------------- summary

def apply_summary(doc, summary: str) -> None:
    """The summary is the last non-empty normal paragraph before the first Heading-1."""
    target = None
    for p in _paras(doc):
        if _is_heading(p):
            break
        if _text(p):
            target = p
    if target is None:
        raise RuntimeError("could not locate summary paragraph")
    _set_single_run_text(target, summary)


# ----------------------------------------------------------------------------- skills

def _skills_block(doc) -> list[Paragraph]:
    """Normal paragraphs between the TECHNICAL SKILLS heading and the next heading."""
    out, inside = [], False
    for p in _paras(doc):
        if _is_heading(p):
            if _text(p).upper().startswith("TECHNICAL SKILLS"):
                inside = True
                continue
            if inside:
                break
        elif inside and _text(p):
            out.append(p)
    return out


def _write_skill_line(p: Paragraph, label: str, value: str) -> None:
    """Rewrite a skills paragraph as [bold label][plain value], reusing existing run formatting.

    The base lines are either "Label: value" (space after colon) or "Label:value" (no space, e.g.
    "AI/ML:"). We keep a trailing space after the colon for readability except mirror the base's
    tightest form only when the label already had none — simplest: always "Label: ".
    """
    runs = p.runs
    if len(runs) < 2:
        # ensure we have a bold-label run and a plain-value run to inherit formatting from
        while len(p.runs) < 2:
            p.add_run("")
        runs = p.runs
    label_run, value_run = runs[0], runs[1]
    label_run.text = f"{label}: "
    label_run.bold = True
    value_run.text = value
    value_run.bold = False
    for r in runs[2:]:
        r._r.getparent().remove(r._r)


def apply_skills(doc, skills: list[dict]) -> None:
    block = _skills_block(doc)
    if not block:
        raise RuntimeError("could not locate skills block")
    template = block[0]._p
    anchor = block[-1]._p  # insert clones after the last existing line, then drop originals

    # Build the new lines by cloning the template paragraph (inherits style/spacing/rPr).
    prev = anchor
    new_ps = []
    for item in skills:
        clone = copy.deepcopy(template)
        prev.addnext(clone)
        prev = clone
        new_ps.append(clone)

    for clone, item in zip(new_ps, skills):
        _write_skill_line(Paragraph(clone, template.getparent()), item["label"], item["value"])

    for p in block:  # remove the originals
        p._p.getparent().remove(p._p)


# ----------------------------------------------------------------------------- experience bullets

def _section_bullets(doc, match: str) -> tuple[Paragraph, list[Paragraph]]:
    """Find the job header containing `match` and return (header, [its bullet paragraphs])."""
    paras = _paras(doc)
    header_idx = None
    for i, p in enumerate(paras):
        if not _is_heading(p) and not _is_bullet(p) and match.lower() in _text(p).lower() and p.runs and p.runs[0].bold:
            header_idx = i
            break
    if header_idx is None:
        raise RuntimeError(f"experience section not found: {match!r}")
    bullets = []
    for p in paras[header_idx + 1:]:
        if _is_heading(p):
            break
        if _is_bullet(p):
            bullets.append(p)
        elif bullets:  # a non-bullet line after bullets started => next section header
            break
    return paras[header_idx], bullets


def _has_hyperlink(p: Paragraph) -> bool:
    return len(p._p.findall(qn("w:hyperlink"))) > 0


def _find_subline_template(doc):
    """A clean date/subline paragraph (plain, no hyperlink) to clone for a new entry's sub-line.

    The flagship (Reazy) sub-line carries live hyperlinks, so we borrow a plain one (e.g. the Economist
    date line) as the template to avoid copying stray links into the new entry.
    """
    paras = _paras(doc)
    seen_heading = False  # skip the centered name/subtitle header block, which precedes any Heading-1
    for i, p in enumerate(paras):
        if _is_heading(p):
            seen_heading = True
            continue
        is_title = seen_heading and not _is_bullet(p) and p.runs and p.runs[0].bold and _text(p)
        if is_title and i + 1 < len(paras):
            nxt = paras[i + 1]
            if (not _is_heading(nxt) and not _is_bullet(nxt) and not _has_hyperlink(nxt)
                    and _text(nxt) and not (nxt.runs and nxt.runs[0].bold)):
                return nxt._p
    raise RuntimeError("could not find a clean sub-line template to clone")


def insert_experience(doc, before: str, heading: str, subline: str, bullets: list[str]) -> None:
    """Insert a new experience entry (bold heading + plain sub-line + bullets) directly above `before`."""
    target_header, tgt_bullets = _section_bullets(doc, before)
    if not tgt_bullets:
        raise RuntimeError(f"need a bullet in {before!r} to clone list formatting")
    parent = target_header._p.getparent()
    anchor = target_header._p

    new_header = copy.deepcopy(target_header._p)      # bold job-title formatting
    new_subline = copy.deepcopy(_find_subline_template(doc))  # plain grey date line
    new_bullets = [copy.deepcopy(tgt_bullets[0]._p) for _ in bullets]  # numbered list items

    for el in [new_header, new_subline, *new_bullets]:
        anchor.addprevious(el)

    _set_single_run_text(Paragraph(new_header, parent), heading)
    _set_single_run_text(Paragraph(new_subline, parent), subline)
    for el, text in zip(new_bullets, bullets):
        _set_single_run_text(Paragraph(el, parent), text)


def apply_experience_title(doc, match: str, title: str) -> None:
    """Override an experience entry's bold header line (e.g. add/drop a '(Founder)' parenthetical).

    The header is matched by `match` and carries no hyperlinks (those live on the sub-line below it),
    so a single-run text swap is safe.
    """
    header, _ = _section_bullets(doc, match)
    _set_single_run_text(header, title)


def apply_experience(doc, match: str, bullets: list[str]) -> None:
    _, existing = _section_bullets(doc, match)
    if not existing:
        raise RuntimeError(f"no bullets to template from in section: {match!r}")
    template = existing[0]._p
    parent = template.getparent()

    prev = existing[-1]._p
    clones = []
    for _ in bullets:
        clone = copy.deepcopy(template)
        prev.addnext(clone)
        prev = clone
        clones.append(clone)

    for clone, text in zip(clones, bullets):
        _set_single_run_text(Paragraph(clone, parent), text)

    for p in existing:  # drop the originals
        p._p.getparent().remove(p._p)


# ----------------------------------------------------------------------------- driver

def apply_tighten(doc, t: dict) -> None:
    """Compact vertical spacing to reclaim lines (used before cutting content to hold one page).

    Keys (all optional): margin_tb_in, heading_before_pt, heading_after_pt, entry_before_pt,
    bullet_after_pt (caps space-after on list items), line_spacing.
    """
    from docx.shared import Inches, Pt
    if "margin_tb_in" in t:
        for s in doc.sections:
            s.top_margin = Inches(t["margin_tb_in"])
            s.bottom_margin = Inches(t["margin_tb_in"])
    for p in _paras(doc):
        pf = p.paragraph_format
        if _is_heading(p):
            if "heading_before_pt" in t:
                pf.space_before = Pt(t["heading_before_pt"])
            if "heading_after_pt" in t:
                pf.space_after = Pt(t["heading_after_pt"])
            continue
        if "line_spacing" in t:
            pf.line_spacing = t["line_spacing"]
        is_entry_header = bool(p.runs) and p.runs[0].bold and not _is_bullet(p) and _text(p)
        if is_entry_header and "entry_before_pt" in t:
            pf.space_before = Pt(t["entry_before_pt"])
        if _is_bullet(p) and "bullet_after_pt" in t:
            pf.space_after = Pt(t["bullet_after_pt"])
        elif not _is_bullet(p) and not is_entry_header and "normal_after_pt" in t:
            pf.space_after = Pt(t["normal_after_pt"])


def render(base: Path, plan: dict, out: Path) -> None:
    if not base.exists():
        raise SystemExit(
            f"base résumé not found: {base}\n"
            f"  /tailor-cv edits YOUR résumé in place — it needs one to start from.\n"
            f"  Put your résumé there (`cp your-resume.docx {base}`), or run /setup which copies it for\n"
            f"  you. config/example/cv-base.docx is a worked example of the structure the plan edits.")
    doc = Document(str(base))
    if plan.get("title"):
        apply_title(doc, plan["title"])
    if plan.get("tagline"):
        apply_tagline(doc, plan["tagline"])
    if plan.get("summary"):
        apply_summary(doc, plan["summary"])
    if plan.get("skills"):
        apply_skills(doc, plan["skills"])
    for section in plan.get("experience", []):
        if section.get("title"):
            apply_experience_title(doc, section["match"], section["title"])
        if section.get("bullets"):
            apply_experience(doc, section["match"], section["bullets"])
    for entry in plan.get("insert_experience", []):
        insert_experience(doc, entry["before"], entry["heading"], entry["subline"], entry["bullets"])
    if plan.get("tighten"):
        apply_tighten(doc, plan["tighten"])
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def to_pdf(docx_path: Path) -> Path:
    """Convert to PDF via headless LibreOffice for faithful Word rendering."""
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not found; install it or omit --pdf")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=True, capture_output=True, text=True,
    )
    pdf = docx_path.with_suffix(".pdf")
    if not pdf.exists():
        raise RuntimeError("PDF conversion did not produce a file")
    return pdf


def _find_soffice() -> str | None:
    import shutil
    for c in ("soffice", "libreoffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice"):
        if shutil.which(c) or Path(c).exists():
            return shutil.which(c) or c
    return None


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--base", required=True, type=Path)
    ap.add_argument("--plan", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    ap.add_argument("--pdf", action="store_true", help="also emit a PDF next to the docx")
    args = ap.parse_args()

    plan = json.loads(args.plan.read_text())
    render(args.base, plan, args.out)
    print(f"wrote {args.out}")
    if args.pdf:
        pdf = to_pdf(args.out)
        print(f"wrote {pdf}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
