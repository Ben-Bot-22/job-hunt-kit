#!/usr/bin/env python3
"""Generate a formatted cover-letter .docx (matches the résumé: Arial, centered name header).

Usage: make_cover_letter.py --letter <letter.json> --out <path.docx> [--pdf]

**This file is machinery and holds no content**, the same split `render_cv.py` already makes: the
base document and the edit-plan are yours and live outside this directory, and the script only knows
how to lay them out. Two inputs, for two different lifetimes:

  * the **header** — your name, your title, your contact line — from `profile/profile.yaml →
    identity`, because it is the same on every letter you will ever send;
  * the **letter** — date, salutation, body paragraphs, closing — from a JSON file, because it is
    different every time. `--letter` defaults to `profile/cover-letter.json` (your standing letter);
    point it at one inside `applications/<folder>/` for a tailored one.

Letter schema, every key optional except `body`:

  {"date": "July 6, 2026", "salutation": "Dear Hiring Manager,",
   "body": ["first paragraph", "second paragraph"],
   "closing": "Thank you & best of luck with the search!"}

It used to carry one person's name, address and three paragraphs of their biography as module
constants, which is how a shipping file comes to contain a résumé.
"""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

# Run directly (`python cv/scripts/make_cover_letter.py …`), sys.path[0] is this directory, so
# `core` is not importable without help. The skill and the README both invoke it that way.
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from core.settings import PROFILE_DIR, load as _load_yaml

#: Where the standing letter lives when `--letter` is not given. `PROFILE_DIR` rather than a literal
#: `profile/`, so `JOBSDB_CONFIG_HOME` moves this with the rest of the configuration.
DEFAULT_LETTER = PROFILE_DIR / "cover-letter.json"

DEFAULT_SALUTATION = "Dear Hiring Manager,"
DEFAULT_CLOSING = "Thank you & best of luck with the search!"
GREY = RGBColor(0x55, 0x55, 0x55)
LINK = RGBColor(0x1A, 0x0D, 0xAB)   # the blue the résumé's live links use


def identity() -> tuple[str, str, str]:
    """Name, title and contact line from `profile/profile.yaml → identity`.

    Missing is loud rather than defaulted: a cover letter with a blank name header is a document you
    send before you notice, which is the expensive direction.
    """
    block = (_load_yaml(PROFILE_DIR / "profile.yaml").get("identity") or {})
    name = str(block.get("name", "")).strip()
    if not name:
        raise SystemExit(
            f"no `identity: name:` in {PROFILE_DIR / 'profile.yaml'} — a cover letter needs a name "
            "on it. See config/example/profile.yaml for the shape."
        )
    return name, str(block.get("title", "")).strip(), str(block.get("contact", "")).strip()


def letter(path: Path) -> dict:
    """The per-letter content. A missing file names the path and the schema rather than the KeyError."""
    if not path.is_file():
        raise SystemExit(
            f"no letter at {path}. Write one — "
            '{"date": ..., "salutation": ..., "body": ["para", ...], "closing": ...} — '
            "or pass --letter."
        )
    data = json.loads(path.read_text(encoding="utf-8"))
    if not data.get("body"):
        raise SystemExit(f"{path} has no `body` paragraphs, so there is no letter to render.")
    return data


def _add_hyperlink(p, url: str, text: str, *, size: float):
    """Append a real, clickable hyperlink run to paragraph `p`.

    python-docx has no first-class hyperlink API, so we add the relationship by hand and wrap a styled
    run in a `w:hyperlink` element. Same mechanism the résumé's base docx uses for its live links — this
    is what lets a cover letter carry active links (portfolio, video) rather than plain text a reader
    has to retype.
    """
    part = p.part
    r_id = part.relate_to(url,
                          "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "1A0DAB"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    rfonts = OxmlElement("w:rFonts"); rfonts.set(qn("w:ascii"), "Arial"); rfonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    wt = OxmlElement("w:t"); wt.text = text; run.append(wt)
    link.append(run)
    p._p.append(link)
    return link


def _render_contact(p, contact: str, size: float = 10):
    """Render the `a | b | c` contact line, making the website token a live link.

    The website is detected, not hardcoded: the last `|`-separated segment that looks like a bare
    domain (has a dot, no `@`, no spaces) is linkified to `https://<that>`. So an email segment stays
    plain and whoever owns the profile gets their own site linked — nothing personal is spelled here.
    """
    parts = [seg.strip() for seg in re.split(r"\s*\|\s*", contact) if seg.strip()]
    site_idx = None
    for i, seg in enumerate(parts):
        if "@" not in seg and " " not in seg and re.fullmatch(r"[\w-]+(?:\.[\w-]+)+", seg):
            site_idx = i
    if site_idx is None:
        _run(p, contact, size=size, color=GREY); return
    for i, seg in enumerate(parts):
        if i:
            _run(p, "  |  ", size=size, color=GREY)
        if i == site_idx:
            _add_hyperlink(p, "https://" + seg, seg, size=size)
        else:
            _run(p, seg, size=size, color=GREY)


def _run(p, text, *, size, bold=False, color=None):
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return r


def _bottom_border(p):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    ppr.append(pbdr)


def build(out: Path, content: dict) -> None:
    name, title, contact = identity()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # --- centered header (matches résumé) ---
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_after = Pt(1)
    _run(h, name, size=18, bold=True)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(1)
    _run(t, title, size=12, color=GREY)

    # Contact line, with its website token rendered as a live link (the résumé does the same) —
    # nothing added to the header, just the domain already there made clickable.
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    _render_contact(c, contact)
    _bottom_border(c)

    # --- date ---
    d = doc.add_paragraph()
    d.paragraph_format.space_before = Pt(6)
    d.paragraph_format.space_after = Pt(12)
    _run(d, str(content.get("date", "")), size=11, color=GREY)

    # --- salutation ---
    s = doc.add_paragraph()
    s.paragraph_format.space_after = Pt(10)
    _run(s, str(content.get("salutation") or DEFAULT_SALUTATION), size=11)

    # --- body ---
    # A paragraph is a plain string, OR a list of segments for inline links: a string segment is plain
    # text, a {"text","url"} segment is a clickable run. Lets the 3-min-video link live inside a sentence
    # rather than as a header ornament.
    for para in content["body"]:
        b = doc.add_paragraph()
        b.paragraph_format.space_after = Pt(10)
        b.paragraph_format.line_spacing = 1.15
        if isinstance(para, list):
            for seg in para:
                if isinstance(seg, dict):
                    _add_hyperlink(b, str(seg["url"]), str(seg["text"]), size=11)
                else:
                    _run(b, str(seg), size=11)
        else:
            _run(b, str(para), size=11)

    # --- closing ---
    cl = doc.add_paragraph()
    cl.paragraph_format.space_before = Pt(6)
    cl.paragraph_format.space_after = Pt(2)
    _run(cl, str(content.get("closing") or DEFAULT_CLOSING), size=11)

    sign = doc.add_paragraph()
    _run(sign, name, size=11, bold=True)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def to_pdf(docx_path: Path) -> Path:
    import shutil
    soffice = shutil.which("soffice") or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                    str(docx_path.parent), str(docx_path)], check=True, capture_output=True, text=True)
    return docx_path.with_suffix(".pdf")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out", required=True, type=Path)
    ap.add_argument("--letter", type=Path, default=DEFAULT_LETTER,
                    help=f"the letter's content as JSON (default: {DEFAULT_LETTER})")
    ap.add_argument("--pdf", action="store_true")
    args = ap.parse_args()
    build(args.out, letter(args.letter))
    print(f"wrote {args.out}")
    if args.pdf:
        print(f"wrote {to_pdf(args.out)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
