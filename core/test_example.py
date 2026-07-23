"""The shipped example: a fictional job seeker who is also the test fixture.
Run:  .venv/bin/python -m pytest core/ -q

`config/example/` exists so a stranger's first run is a real run — clone, point one environment
variable at it, get scored jobs — and it is the test fixture for exactly the same reason the ticket
gave: a demo nothing exercises is a demo that rots, and the day it rots is a stranger's first five
minutes.

Everything here is offline. The failure directions, in the order they cost something:

  * **A real person's details in a shipped, tracked, about-to-be-published file.** The example is what
    replaces the repo owner's own profile when this goes out. An email address or a Google Sheet id
    left in it is not a bug you find later — it is published.
  * **An example that no longer loads.** It goes through the same Pydantic schema as the real settings
    file, so a key renamed in `core/settings.py` and updated only in `config/settings.yaml` fails here
    rather than in a stranger's terminal.
  * **The example drifting out of shape from the real config.** Identity is a fixed set of keys; an
    example missing one teaches the wrong shape to the next person who copies it, and to `/setup`.
  * **Seeding overwriting a configured profile.** `profile/rubric.md` is the most-tuned file in the
    repo and there is no undo. `seed()` must skip what is already there, by default, always.
"""
from __future__ import annotations

from pathlib import Path

import pytest
import yaml

from .example import DESTINATIONS, EXAMPLE_DIR, seed
from .settings import (CONFIG_HOME_ENV, PROFILE_DIR, REPO_ROOT, SETTINGS_PATH, Settings, profile_exists,
                       is_example_profile, load, validate)

#: Two tests below diff the example against the REPO OWNER's real `profile/profile.yaml`. That file is
#: the one thing the public snapshot never ships, so on any clone but the owner's there is nothing to
#: diff against and the comparison has no subject. Skipping says so; erroring would tell a reader who
#: just cloned the repo that it is broken. See docs/agents/tests.md.
needs_owner_profile = pytest.mark.skipif(
    not profile_exists(),
    reason="diffs the example against the owner's profile/profile.yaml, which the public tree omits "
           "by design. Seed a profile with `python -m core.example` to run the rest of the suite.")

EXAMPLE_SETTINGS = EXAMPLE_DIR / "settings.yaml"
EXAMPLE_PROFILE = EXAMPLE_DIR / "profile.yaml"
EXAMPLE_RUBRIC = EXAMPLE_DIR / "rubric.md"


def _scannable_text(path: Path) -> str:
    """The text of an example file for the leak/secret scans — including inside the `.docx` base résumé.

    Most example files are text; `cv-base.docx` is a zip, so `read_text` raises. The guard still has to
    see its words (a leaked name inside a Word document is still a leak), so a `.docx` is unzipped and
    its paragraph text returned. Any other binary that can't be read is skipped with an empty string.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        if path.suffix == ".docx":
            from docx import Document
            return "\n".join(p.text for p in Document(str(path)).paragraphs)
        return ""


# --------------------------------------------------------------------------------------------------
# It ships, and it is complete
# --------------------------------------------------------------------------------------------------

def test_the_example_ships_every_file_a_configuration_needs() -> None:
    """A partial example is worse than none: it runs, and then fails on the file nobody copied.

    The list is `DESTINATIONS`, so adding a config file without adding it to the example fails here.
    """
    for name in DESTINATIONS:
        assert (EXAMPLE_DIR / name).exists(), f"config/example/{name} is missing"
    assert (EXAMPLE_DIR / "README.md").exists()


def test_the_example_settings_validate_against_the_real_schema() -> None:
    """The same gate the user's own file goes through — `extra="forbid"` and every bound.

    Failure direction: a settings key renamed in `core/settings.py` and updated only in
    `config/settings.yaml`. The example then fails on a stranger's first command, which is the one
    moment in this project's life where nobody has anyone to ask.
    """
    data = load(EXAMPLE_SETTINGS)
    assert validate(data, EXAMPLE_SETTINGS) == data


def test_the_example_names_a_value_for_every_top_level_setting() -> None:
    """Complete, not minimal. The example is read as documentation as much as it is run.

    A stranger who cannot see `liveness:` in the file they were given does not learn it exists; the
    schema is the reference, but nobody reads the reference first.
    """
    data = load(EXAMPLE_SETTINGS)
    assert set(data) == set(Settings.model_fields)


@needs_owner_profile
def test_the_example_identity_has_exactly_the_shape_of_the_real_one() -> None:
    """Same keys as `profile/profile.yaml`, so the example teaches the true shape.

    Failure direction: a field added to the repo owner's profile and not to the example. `/setup`
    seeds from the example, so the new field silently stops existing for every new user.
    """
    example = yaml.safe_load(EXAMPLE_PROFILE.read_text())
    real = yaml.safe_load((REPO_ROOT / "profile" / "profile.yaml").read_text())
    assert set(example) == set(real)


def test_the_channels_a_stranger_cannot_use_are_off_and_the_rest_are_on() -> None:
    """The first-run posture, asserted rather than described in a comment.

    `mail` needs macOS and a configured Apple Mail account; `gmail` is a stub that raises. Both on
    would make the shipped demo fail for everyone not sitting at the repo owner's desk. `paste` and
    `boards` need no key and no OAuth, which is the promise the example exists to keep.
    """
    channels = load(EXAMPLE_SETTINGS)["channels"]
    assert channels["mail"]["enabled"] is False
    assert channels["gmail"]["enabled"] is False
    assert channels["paste"]["enabled"] is True
    assert channels["boards"]["enabled"] is True
    # Named boards, not an empty watchlist: with nothing to paste, this is the only thing that makes a
    # first run produce jobs at all. All public, keyless and real, and leaning high-volume so a quiet
    # week still returns jobs rather than reading as broken — the counts live in starter-boards.md.
    assert channels["boards"]["greenhouse"] == ["stripe", "databricks", "mongodb",
                                                "gitlab", "cloudflare", "anthropic"]
    # gopuff, not leverdemo: the demo board's postings are old and report 0 in-window, which looks broken.
    assert channels["boards"]["lever"] == ["gopuff"]


# --------------------------------------------------------------------------------------------------
# Unmistakably fictional
# --------------------------------------------------------------------------------------------------

@pytest.mark.skipif(
    not profile_exists() or is_example_profile(),
    reason="reads the owner's real profile to diff against the example; on a clone profile IS the "
           "example, or has not been seeded at all")
def test_no_real_person_is_in_the_example() -> None:
    """The example replaces the repo owner's profile when this repo is published. It cannot carry him.

    Checked by shape as well as by name: `example.invalid` is reserved by RFC 2606 and can never be
    registered, so an address there is provably nobody's. The Sheet id is empty, because a Drive id is
    the one value here that is a capability rather than a fact.

    **Every denied value is read out of the real profile, never spelled here.** This file ships, and a
    guard that protects the owner's surname by writing it down is a guard that leaks the thing it
    guards — `scripts/test_leaks.py` used to find this exact line. It also means the test works for
    whoever owns the clone, rather than only for the person who wrote it.
    """
    real_profile = yaml.safe_load((REPO_ROOT / "profile" / "profile.yaml").read_text())
    real_account = str(real_profile.get("inbox", {}).get("account", "")).lower()
    real_sheet = str(real_profile.get("applied_sheet", ""))
    real_name = str((real_profile.get("identity") or {}).get("name", "")).strip().lower()
    surname = real_name.split()[-1] if real_name else ""

    for path in EXAMPLE_DIR.rglob("*"):
        if not path.is_file():
            continue
        text = _scannable_text(path)
        low = text.lower()
        assert real_account not in low, f"{path.name} carries the repo owner's inbox address"
        assert real_sheet and real_sheet not in text, f"{path.name} carries a real Google Sheet id"
        assert surname and surname not in low, f"{path.name} names the repo owner"

    example = yaml.safe_load(EXAMPLE_PROFILE.read_text())
    assert example["inbox"]["account"].endswith("@example.invalid")
    assert example["applied_sheet"] == ""


def test_the_example_base_resume_has_the_anchors_render_keys_on() -> None:
    """`cv-base.docx` is the worked example `/tailor-cv` points a stranger at. render_cv.py matches by
    anchor text, not index, so the example must carry them or the example itself is broken: the three
    Heading-1 sections, a bold job header, and at least one INLINE-numbered bullet (the style-level
    numbering on 'List Bullet' is not what `_is_bullet` looks for). Guards the base against silent rot."""
    from docx import Document
    from docx.oxml.ns import qn

    base = EXAMPLE_DIR / "cv-base.docx"
    assert base.is_file(), "the example base résumé is missing"
    doc = Document(str(base))

    headings = {p.text.strip().upper() for p in doc.paragraphs if p.style.name.startswith("Heading")}
    for anchor in ("TECHNICAL SKILLS", "EXPERIENCE", "EDUCATION"):
        assert anchor in headings, f"cv-base.docx has no {anchor!r} heading for render_cv to anchor on"

    def _inline_numbered(p) -> bool:
        ppr = p._p.find(qn("w:pPr"))
        return ppr is not None and ppr.find(qn("w:numPr")) is not None

    assert any(_inline_numbered(p) for p in doc.paragraphs), "no inline-numbered bullet — render_cv won't see any"
    assert any(p._p.findall(qn("w:hyperlink")) for p in doc.paragraphs), "no hyperlink line (the contact anchor)"


def test_no_secret_reaches_the_example() -> None:
    """Same check `triage/test_config.py` runs over the real config files, for the same reason.

    What gets committed by accident is a value somebody pasted while debugging, not a field somebody
    designed — so this looks for the shape of a key rather than for a key's name.
    """
    for path in EXAMPLE_DIR.rglob("*"):
        if path.is_file():
            text = _scannable_text(path)
            assert "sk-ant-" not in text and "sk-proj-" not in text


def test_the_example_rubric_is_a_different_person_not_a_redacted_one() -> None:
    """A sanitised copy of the owner's rubric would teach that the rubric is a template with names in it.

    Robin PREFERS permanent salaried work at a stated salary floor and will take a contract that clears
    an hourly one; the repo owner is contract-first at a different hourly floor, in a different city, on
    a different stack. The two anchors still disagree about the most basic question the file asks, which
    is the demonstration: this file is prose, and it is *yours*.
    """
    rubric = EXAMPLE_RUBRIC.read_text(encoding="utf-8")
    assert "PERMANENT, salaried" in rubric and "$115k HARD FLOOR" in rubric
    assert "$50/hr" not in rubric and "Reazy" not in rubric
    # Same load-bearing sections as the real anchor, so the example teaches the shape that works.
    for heading in ("IDEAL ROLE", "DOWNRANK", "SCORING DISCIPLINE", "CALIBRATION",
                    "HARD FILTERS", "CANDIDATE"):
        assert heading in rubric
    # No front matter and no title: the whole file is the prompt, and a `# Rubric` heading here would
    # be read by the model as an instruction. Pinned in the example because this is the file people
    # copy, and a markdown title is exactly what a helpful editor adds to it.
    assert rubric.startswith("IDEAL ROLE ")


# --------------------------------------------------------------------------------------------------
# `JOBSDB_CONFIG_HOME` — how a fresh clone runs against it
# --------------------------------------------------------------------------------------------------

def test_the_config_home_variable_moves_both_halves_together(monkeypatch) -> None:
    """One variable, both halves. Settings and identity have to move together or not at all.

    Failure direction if they could move apart: a demo run scored against the shipped example's rubric
    but reading the *real* inbox account, which is the one combination that reaches a real mailbox
    while claiming to be an example.
    """
    from . import settings as settings_module
    monkeypatch.setenv(CONFIG_HOME_ENV, str(EXAMPLE_DIR))
    home = settings_module.config_home()
    assert home == EXAMPLE_DIR
    assert (home / "settings.yaml").exists() and (home / "profile.yaml").exists()
    assert (home / "rubric.md").exists() and (home / "skiplist.md").exists()


def test_unset_means_the_normal_two_place_layout() -> None:
    """The default is what every previous run used: `profile/` and `config/settings.yaml`.

    Asserted because this ticket moved both constants behind a resolver, and a resolver that silently
    answered `config/example/` would score the repo owner's real inbox against a fictional rubric.
    """
    assert PROFILE_DIR == REPO_ROOT / "profile"
    assert SETTINGS_PATH == REPO_ROOT / "config" / "settings.yaml"


# --------------------------------------------------------------------------------------------------
# Seeding — what `/setup` calls
# --------------------------------------------------------------------------------------------------

def test_seeding_writes_the_whole_configuration(tmp_path: Path) -> None:
    dest = {name: tmp_path / target.parent.name / name for name, target in DESTINATIONS.items()}
    written, skipped = seed(destinations=dest)

    assert skipped == []
    assert sorted(p.name for p in written) == sorted(DESTINATIONS)
    for name, target in dest.items():
        assert target.read_text() == (EXAMPLE_DIR / name).read_text()


def test_seeding_never_overwrites_what_is_already_there(tmp_path: Path) -> None:
    """**The one behaviour that cannot be got wrong.**

    `profile/rubric.md` is a month of calibration and there is no undo. A second `/setup` on a
    configured clone — which is exactly what a user runs when they want to change one answer — must
    report what it left alone rather than restore factory settings.
    """
    dest = {name: tmp_path / name for name in DESTINATIONS}
    tuned = "IDEAL ROLE: the one I spent a month tuning\n"
    dest["rubric.md"].write_text(tuned)

    written, skipped = seed(destinations=dest)

    assert dest["rubric.md"].read_text() == tuned
    assert skipped == [dest["rubric.md"]]
    assert dest["rubric.md"] not in written
    assert dest["profile.yaml"].exists()      # the rest still landed


def test_force_is_how_you_overwrite_and_it_has_to_be_asked_for(tmp_path: Path) -> None:
    dest = {name: tmp_path / name for name in DESTINATIONS}
    dest["rubric.md"].write_text("mine\n")

    written, skipped = seed(destinations=dest, force=True)

    assert skipped == []
    assert dest["rubric.md"].read_text() == (EXAMPLE_DIR / "rubric.md").read_text()
    assert len(written) == len(DESTINATIONS)


def test_a_missing_example_file_is_loud(tmp_path: Path) -> None:
    """Half a seed is worse than none — a profile with no rubric fails three phases later, at the
    analyzer, with a message about a missing anchor rather than about a missing copy."""
    (tmp_path / "profile.yaml").write_text("inbox: {}\n")
    with pytest.raises(FileNotFoundError, match="rubric.md"):
        seed(example_dir=tmp_path, destinations={n: tmp_path / "out" / n for n in DESTINATIONS})


def test_the_default_seed_targets_are_where_the_tool_actually_reads() -> None:
    """`python -m core.example` on a fresh clone has to land the config where the tool looks for it.

    Asserted against the paths rather than by running `seed()` with its defaults: this suite runs on a
    configured clone, and a test that writes into `profile/` to prove it doesn't is not a trade worth
    making. That every one of them already exists here is also why the real command is a no-op.
    """
    assert DESTINATIONS["settings.yaml"] == SETTINGS_PATH
    assert DESTINATIONS["rubric.md"] == PROFILE_DIR / "rubric.md"
    assert DESTINATIONS["profile.yaml"] == PROFILE_DIR / "profile.yaml"
    assert all(p.exists() for p in DESTINATIONS.values())
